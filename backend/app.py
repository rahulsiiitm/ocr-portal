from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pytesseract
from PIL import Image
import io
from docx import Document
import Levenshtein
import jellyfish
import os
import re

# -----------------------------
# Tesseract path
# -----------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

app = Flask(__name__)
CORS(app)

# ---------------------------------------------------
# LOAD HINDI DICTIONARY
# ---------------------------------------------------

def load_hindi_dictionary():
    hindi_words = set()
    file_path = "hindi.txt"

    if not os.path.exists(file_path):
        print("ERROR: hindi.txt not found. Using fallback dictionary.")
        return {"हिंदी", "भारत", "देवनागरी", "में", "है", "लिखी", "जाती"}

    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            # 1. Remove metadata
            line = re.sub(r"\\", "", line)
            line = re.sub(r"\(.*?\)", "", line)
            
            # 2. Extract Devanagari word sequences 
            words_found = re.findall(r"[\u0900-\u097F]+", line)
            
            for w in words_found:
                if len(w) > 1:
                    hindi_words.add(w)

    print(f"Dictionary Loaded: {len(hindi_words)} words.")
    return hindi_words

HINDI_DICTIONARY = load_hindi_dictionary()

# ---------------------------------------------------
# CLEAN TOKEN
# ---------------------------------------------------

def clean_token(word):
    """
    Strips punctuation but keeps characters. 
    Using a regex to delete non-Hindi characters is dangerous 
    because your OCR uses 'hin+eng' and would delete valid English words.
    """
    return word.strip('.,!?:;|।"\'()[]{}')

# ---------------------------------------------------
# DISTANCE CALCULATIONS
# ---------------------------------------------------

def get_distances(word, benchmark):
    # Hamming Distance
    if len(word) == len(benchmark):
        hamming = sum(c1 != c2 for c1, c2 in zip(word, benchmark))
    else:
        hamming = "N/A"

    # LCS Distance (Dynamic Programming)
    m, n = len(word), len(benchmark)
    dp = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if word[i - 1] == benchmark[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
            else:
                dp[i][j] = max(dp[i - 1][j], dp[i][j - 1])

    lcs_len = dp[m][n]
    lcs_dist = (m + n) - 2 * lcs_len

    return {
        "hamming": hamming,
        "lcs": lcs_dist,
        "levenshtein": Levenshtein.distance(word, benchmark),
        "jaro": round(jellyfish.jaro_similarity(word, benchmark), 3)
    }

# ---------------------------------------------------
# SPELL CHECK
# ---------------------------------------------------

@app.route("/check-spelling", methods=["POST"])
def check_spelling():
    data = request.json
    text = data.get("text", "")
    tokens = text.split()
    results = []

    search_pool = list(HINDI_DICTIONARY)[:2000]

    for token in tokens:
        clean_word = clean_token(token)
        if not clean_word:
            continue

        # Check if it's a known Hindi word
        is_hindi = any('\u0900' <= c <= '\u097F' for c in clean_word)
        
        if is_hindi:
            status = "Correct" if clean_word in HINDI_DICTIONARY else "Wrong"
        else:
            status = "Correct"

        partial = None

        if status == "Wrong" and is_hindi:
            for i in range(len(clean_word) - 1, 2, -1):
                prefix = clean_word[:i]
                if prefix in HINDI_DICTIONARY:
                    status = "Partial"
                    partial = f"Correct: {prefix} | Error: {clean_word[i:]}"
                    break

        # Benchmark Selection
        if status == "Correct":
            benchmark = clean_word
        else:
            benchmark = min(search_pool, key=lambda x: Levenshtein.distance(clean_word, x))

        dist = get_distances(clean_word, benchmark)

        results.append({
            "word": token,
            "benchmark": benchmark,
            "status": status,
            "partial": partial,
            **dist
        })

    return jsonify(results)

# ---------------------------------------------------
# OCR IMAGE PROCESSING
# ---------------------------------------------------

@app.route("/process-image", methods=["POST"])
def process_image():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    image = Image.open(file.stream)

    extracted_text = pytesseract.image_to_string(image, lang="hin+eng")

    extracted_text = re.sub(r"\s+", " ", extracted_text).strip()

    words = extracted_text.split()

    stats = {
        "word_count": len(words),
        "sentence_count": len([s for s in extracted_text.split("।") if s.strip()]),
        "char_count": len(extracted_text),
        "avg_word_length": round(sum(len(w) for w in words) / len(words), 2) if words else 0
    }

    return jsonify({
        "text": extracted_text,
        "stats": stats
    })

# ---------------------------------------------------
# DOWNLOAD DOCX
# ---------------------------------------------------

@app.route("/download-docx", methods=["POST"])
def download_docx():
    text = request.json.get("text", "")
    doc = Document()
    doc.add_heading("OCR Extracted Text", 0)
    doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name="extracted_text.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    app.run(debug=True, port=5000)